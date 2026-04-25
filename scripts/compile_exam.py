"""Generación de documentos .docx de examen con permutación anti-fraude.

Funciones principales:
  - build_exam_batch: genera N versiones (Formas) de un examen
  - generate_docx_exam: renderiza un examen en Word con QR y marcadores
  - prepare_exam_shuffling: selecciona preguntas y baraja opciones + orden
"""

from __future__ import annotations

import io
import random
import secrets
import string
from pathlib import Path
from typing import Any

import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from loguru import logger
from PIL import Image

from scripts.supabase_client import (
    create_exam,
    get_answer_key,
    get_questions_by_session,
    save_answer_keys,
)

LETTERS = ["A", "B", "C", "D", "E"]
FORM_LABELS = list(string.ascii_uppercase)  # A, B, C, … Z


# ── Helpers de QR y marcadores fiduciarios ────────────────────────────────────

def _make_qr_image(data: str, box_size: int = 4) -> Image.Image:
    """Genera una imagen PIL de QR con los datos del examen."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=box_size,
        border=2,
    )
    qr.add_data(data)
    qr.make(fit=True)
    return qr.make_image(fill_color="black", back_color="white").convert("RGB")


def _pil_to_bytes(img: Image.Image, fmt: str = "PNG") -> bytes:
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


def _add_fiducial_square(cell, size_px: int = 12) -> None:
    """Inserta un cuadrado negro en una celda de tabla (marcador de esquina)."""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.font.size = Pt(size_px)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.text = "■"


def _add_qr_to_paragraph(paragraph, qr_data: str) -> None:
    """Añade imagen QR inline al párrafo dado."""
    qr_img = _make_qr_image(qr_data)
    img_bytes = _pil_to_bytes(qr_img)
    buf = io.BytesIO(img_bytes)
    run = paragraph.add_run()
    run.add_picture(buf, width=Inches(1.0))


# ── Renderizado de documento ─────────────────────────────────────────────────

def generate_docx_exam(
    exam_id: str,
    exam_code: str,
    subject: str,
    output_path: Path | str,
    form_variant: str = "A",
) -> None:
    """Genera el archivo .docx del examen a partir de las claves en Supabase."""
    output_path = Path(output_path)
    data = get_answer_key(exam_id)

    if not data:
        logger.error(f"Sin preguntas para exam_id={exam_id}")
        return

    doc = Document()

    # Márgenes reducidos para más contenido por página
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Header con marcadores fiduciarios + QR ────────────────────────────
    qr_payload = f"{exam_code}|{form_variant}|{len(data)}"

    # Tabla 3×3 para header (esquina-TL, centro, esquina-TR)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.style = "Table Grid"
    left_cell, mid_cell, right_cell = header_table.rows[0].cells

    _add_fiducial_square(left_cell)
    mid_p = mid_cell.paragraphs[0]
    mid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mid_run = mid_p.add_run(f"Examen de {subject}\nCódigo: {exam_code} | Forma {form_variant}")
    mid_run.bold = True
    mid_run.font.size = Pt(11)

    _add_qr_to_paragraph(right_cell.paragraphs[0], qr_payload)

    doc.add_paragraph("")

    # ── Datos del estudiante ──────────────────────────────────────────────
    doc.add_paragraph("Nombre: _____________________________________________   Código: __________________")
    doc.add_paragraph("Fecha: _______________   Materia: " + subject)
    doc.add_paragraph("─" * 72)

    # ── Cuerpo de preguntas ───────────────────────────────────────────────
    for i, item in enumerate(data, 1):
        q = item["question_bank"]
        shuffled_map: dict[str, str] = item["shuffled_options"]

        orig = {l: q[f"option_{l.lower()}"] for l in LETTERS}

        # Construir mapa: letra_actual → texto_opción
        display: dict[str, str] = {}
        for orig_letter, curr_letter in shuffled_map.items():
            display[curr_letter] = orig[orig_letter]

        p_stem = doc.add_paragraph()
        p_stem.add_run(f"{i}. {q['question_text']}").bold = True

        for letter in LETTERS:
            doc.add_paragraph(f"     {letter}) {display[letter]}", style="Normal")

        doc.add_paragraph("")  # espacio entre preguntas

    # ── Tabla de burbujas ─────────────────────────────────────────────────
    doc.add_page_break()

    # Esquina inferior-izquierda e inferior-derecha (marcadores OMR)
    bottom_table = doc.add_table(rows=1, cols=3)
    bottom_table.style = "Table Grid"
    _add_fiducial_square(bottom_table.rows[0].cells[0])
    _add_fiducial_square(bottom_table.rows[0].cells[2])

    doc.add_heading("Tabla de Respuestas", level=2)
    doc.add_paragraph("Rellene completamente el círculo. Use bolígrafo negro o azul.")

    bubble_table = doc.add_table(rows=len(data) + 1, cols=6)
    bubble_table.style = "Table Grid"

    hdr = bubble_table.rows[0].cells
    hdr[0].text = "Preg."
    for j, letter in enumerate(LETTERS, 1):
        hdr[j].text = letter

    for i, _ in enumerate(data):
        row = bubble_table.rows[i + 1].cells
        row[0].text = str(i + 1)
        for j in range(1, 6):
            row[j].text = "○"  # Círculo Unicode real

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"Examen guardado: {output_path}")


# ── Preparación y barajado ───────────────────────────────────────────────────

def prepare_exam_shuffling(
    instructor_id: str,
    session_id: str,
    exam_code: str,
    subject: str,
    question_count: int = 10,
    form_variant: str = "A",
    course_id: str | None = None,
    parent_exam_id: str | None = None,
    rng: random.Random | None = None,
) -> str:
    """
    Selecciona preguntas aprobadas, baraja opciones Y orden de preguntas,
    persiste las claves y retorna el exam_id.
    """
    if rng is None:
        rng = random.SystemRandom()  # criptográficamente seguro

    pool: list[dict[str, Any]] = get_questions_by_session(session_id, status="aprobada")

    if not pool:
        raise ValueError(f"No hay preguntas aprobadas en la sesión {session_id}")

    if len(pool) < question_count:
        logger.warning(
            f"Se pidieron {question_count} preguntas pero solo hay {len(pool)} aprobadas. "
            "Usando todas."
        )

    selected = rng.sample(pool, min(question_count, len(pool)))
    rng.shuffle(selected)  # permutación del orden de preguntas

    exam = create_exam(
        instructor_id=instructor_id,
        exam_code=exam_code,
        subject=subject,
        total_questions=len(selected),
        course_id=course_id,
        form_variant=form_variant,
        parent_exam_id=parent_exam_id,
    )
    exam_id: str = exam["id"]

    keys_to_save: list[dict[str, Any]] = []
    for order_index, q in enumerate(selected):
        current_letters = LETTERS.copy()
        rng.shuffle(current_letters)
        shuffled_map = {orig: curr for orig, curr in zip(LETTERS, current_letters)}
        mapped_correct = shuffled_map[q["correct_answer"]]

        keys_to_save.append({
            "exam_id": exam_id,
            "question_id": q["id"],
            "order_index": order_index,
            "shuffled_options": shuffled_map,
            "mapped_correct_answer": mapped_correct,
        })

    save_answer_keys(keys_to_save)
    logger.info(
        f"Examen {exam_code} Forma {form_variant}: "
        f"{len(selected)} preguntas barajadas (opciones + orden)."
    )
    return exam_id


# ── Generación en batch (múltiples versiones) ────────────────────────────────

def build_exam_batch(
    instructor_id: str,
    session_id: str,
    base_exam_code: str,
    subject: str,
    question_count: int = 10,
    versions: int = 1,
    output_dir: Path | str = Path("workspace/generated_exams"),
    course_id: str | None = None,
) -> list[str]:
    """
    Genera `versions` versiones del mismo examen (Forma A, B, C…).
    Cada versión tiene orden de preguntas y opciones diferentes.
    Retorna la lista de exam_ids creados.
    """
    output_dir = Path(output_dir)
    exam_ids: list[str] = []
    parent_exam_id: str | None = None

    for i in range(versions):
        form = FORM_LABELS[i % len(FORM_LABELS)]
        exam_code = f"{base_exam_code}-{form}" if versions > 1 else base_exam_code

        rng = random.SystemRandom()

        exam_id = prepare_exam_shuffling(
            instructor_id=instructor_id,
            session_id=session_id,
            exam_code=exam_code,
            subject=subject,
            question_count=question_count,
            form_variant=form,
            course_id=course_id,
            parent_exam_id=parent_exam_id if i > 0 else None,
            rng=rng,
        )

        if i == 0:
            parent_exam_id = exam_id

        output_path = output_dir / f"{exam_code}.docx"
        generate_docx_exam(
            exam_id=exam_id,
            exam_code=exam_code,
            subject=subject,
            output_path=output_path,
            form_variant=form,
        )
        exam_ids.append(exam_id)

    logger.info(f"Batch completo: {len(exam_ids)} examen(es) en {output_dir}")
    return exam_ids


if __name__ == "__main__":
    import sys

    print(
        "Usa el CLI: parciales build-exam --help\n"
        "O invoca directamente:\n"
        "  build_exam_batch(instructor_id, session_id, 'CALC1-2025-001', 'Cálculo I', versions=3)"
    )
